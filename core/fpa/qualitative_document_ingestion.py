from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Optional

import pandas as pd
import PyPDF2
from docx import Document


# ---------------------------------------------------------
# SUPPORTED FILE TYPES
# ---------------------------------------------------------

SUPPORTED_TYPES = {
    ".pdf": "PDF",
    ".docx": "DOCX",
    ".txt": "TXT",
}


# ---------------------------------------------------------
# STRUCTURED DOCUMENT
# ---------------------------------------------------------

@dataclass
class QualitativeDocument:
    filename: str
    file_type: str
    category: str
    source_organization: str
    source_type: str
    authority_status: str
    topic: Optional[str]
    title: Optional[str]
    file_path: str
    page_count: Optional[int]
    character_count: int
    text: str
    ingestion_status: str


# ---------------------------------------------------------
# QUALITATIVE DOCUMENT INGESTION
# ---------------------------------------------------------

class QualitativeDocumentIngestion:

    def __init__(self):
        self.documents = []

    # -----------------------------------------------------
    # FILE VALIDATION
    # -----------------------------------------------------

    def validate_file(self, file_path: str) -> Path:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {file_path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Provided path is not a file: {file_path}"
            )

        extension = path.suffix.lower()

        if extension not in SUPPORTED_TYPES:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {list(SUPPORTED_TYPES.keys())}"
            )

        return path

    # -----------------------------------------------------
    # METADATA VALIDATION
    # -----------------------------------------------------

    def validate_metadata(
        self,
        category: str,
        source_organization: str,
        source_type: str,
        authority_status: str,
    ) -> None:

        if not category or not category.strip():
            raise ValueError(
                "Document category is required."
            )

        if not source_organization or not source_organization.strip():
            raise ValueError(
                "Source organization is required."
            )

        if not source_type or not source_type.strip():
            raise ValueError(
                "Source type is required."
            )

        if not authority_status or not authority_status.strip():
            raise ValueError(
                "Authority status is required."
            )

    # -----------------------------------------------------
    # PDF EXTRACTION
    # -----------------------------------------------------

    def extract_pdf(self, path: Path):

        reader = PyPDF2.PdfReader(str(path))

        pages = []

        for page in reader.pages:
            pages.append(
                page.extract_text() or ""
            )

        text = "\n".join(pages)

        return text, len(reader.pages)

    # -----------------------------------------------------
    # DOCX EXTRACTION
    # -----------------------------------------------------

    def extract_docx(self, path: Path):

        document = Document(str(path))

        paragraphs = []

        for paragraph in document.paragraphs:

            if paragraph.text.strip():
                paragraphs.append(
                    paragraph.text
                )

        text = "\n".join(paragraphs)

        return text, None

    # -----------------------------------------------------
    # TXT EXTRACTION
    # -----------------------------------------------------

    def extract_txt(self, path: Path):

        text = path.read_text(
            encoding="utf-8",
            errors="replace",
        )

        return text, None

    # -----------------------------------------------------
    # TEXT CLEANING
    # -----------------------------------------------------

    def clean_text(self, text: str) -> str:

        lines = []

        for line in text.splitlines():

            # Remove unnecessary whitespace
            line = " ".join(line.split())

            if line:
                lines.append(line)

        return "\n".join(lines)

    # -----------------------------------------------------
    # MAIN INGESTION METHOD
    # -----------------------------------------------------

    def ingest(
        self,
        file_path: str,
        category: str,
        source_organization: str,
        source_type: str,
        authority_status: str,
        topic: Optional[str] = None,
        title: Optional[str] = None,
    ) -> QualitativeDocument:

        # 1. Validate file
        path = self.validate_file(file_path)

        # 2. Validate metadata
        self.validate_metadata(
            category=category,
            source_organization=source_organization,
            source_type=source_type,
            authority_status=authority_status,
        )

        # 3. Identify file type
        extension = path.suffix.lower()

        file_type = SUPPORTED_TYPES[extension]

        # 4. Extract text according to file type
        if extension == ".pdf":

            text, page_count = self.extract_pdf(path)

        elif extension == ".docx":

            text, page_count = self.extract_docx(path)

        elif extension == ".txt":

            text, page_count = self.extract_txt(path)

        else:

            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        # 5. Clean extracted text
        text = self.clean_text(text)

        # 6. Validate extracted content
        if not text.strip():

            status = "failed_empty_text"

        else:

            status = "available"

        # 7. Create structured document
        document = QualitativeDocument(
            filename=path.name,
            file_type=file_type,
            category=category,
            source_organization=source_organization,
            source_type=source_type,
            authority_status=authority_status,
            topic=topic,
            title=title,
            file_path=str(path),
            page_count=page_count,
            character_count=len(text),
            text=text,
            ingestion_status=status,
        )

        # 8. Store document
        self.documents.append(document)

        return document

    # -----------------------------------------------------
    # DATAFRAME OUTPUT
    # -----------------------------------------------------

    def to_dataframe(self) -> pd.DataFrame:

        return pd.DataFrame(
            [
                {
                    key: value
                    for key, value in asdict(document).items()
                    if key != "text"
                }
                for document in self.documents
            ]
        )