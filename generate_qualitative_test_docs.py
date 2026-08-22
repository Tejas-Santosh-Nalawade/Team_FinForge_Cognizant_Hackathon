from pathlib import Path
from docx import Document
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parent

OUTPUT_DIR = (
    PROJECT_ROOT
    / "DATASET"
    / "True_data"
    / "qualitative_corpus"
    / "MD&A"
)

OUTPUT_FILE = OUTPUT_DIR / "MDA_REL_01_Credit_Terms.docx"


def create_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()

    document.add_heading("Management Discussion and Analysis", level=1)

    p = document.add_paragraph()
    p.add_run("Synthetic Test Document – REL_01 Context").bold = True

    document.add_heading("Document Status", level=2)

    p = document.add_paragraph()
    p.add_run("SYNTHETIC TEST DATA").bold = True

    document.add_paragraph(
        "This document has been created solely for testing the "
        "qualitative document ingestion and RAG matching pipeline. "
        "It is not an actual company filing, management report, "
        "board memo, or authoritative accounting document."
    )

    document.add_heading("Document Metadata", level=2)

    metadata = [
        ("Document Type", "MD&A"),
        ("Document Category", "Management Commentary"),
        ("Test Flag", "REL_01"),
        ("Reporting Period", "FY2026"),
        ("Source Type", "Synthetic Test Document"),
        ("Authority Status", "Non-authoritative"),
    ]

    for field, value in metadata:
        p = document.add_paragraph()
        p.add_run(f"{field}: ").bold = True
        p.add_run(value)

    document.add_heading("Financial Performance Context", level=2)

    document.add_paragraph(
        "During FY2026, Accounts Receivable increased by "
        "approximately 38% compared with the prior period, "
        "while Revenue increased by approximately 12%."
    )

    document.add_paragraph(
        "The increase in Accounts Receivable relative to Revenue "
        "reflects changes in customer payment terms associated "
        "with several large multi-year contracts."
    )

    document.add_heading("Management Explanation", level=2)

    document.add_paragraph(
        "During the period, the company extended standard customer "
        "credit terms from 30 days to 60 days for selected large "
        "multi-year contracts."
    )

    document.add_paragraph(
        "The extended payment terms were provided to support the "
        "closing and execution of large multi-year customer "
        "agreements. As a result, certain customers had additional "
        "time to settle outstanding invoices."
    )

    document.add_paragraph(
        "The change in payment terms contributed to higher "
        "outstanding trade receivables at the reporting date. "
        "Management therefore considers the increase in Accounts "
        "Receivable to be primarily related to the commercial "
        "terms associated with these contracts."
    )

    document.add_heading("Working Capital Impact", level=2)

    document.add_paragraph(
        "The increase in trade receivables temporarily increased "
        "working capital requirements because cash collection "
        "occurs later under the extended customer payment terms."
    )

    document.add_paragraph(
        "Management continues to monitor customer collections, "
        "outstanding invoices, payment patterns, and credit "
        "exposure associated with these contracts."
    )

    document.add_heading("Related Analytical Flag", level=2)

    document.add_paragraph("Math Engine Flag: REL_01")

    document.add_paragraph(
        "Flag condition: Accounts Receivable growth of approximately "
        "38% compared with Revenue growth of approximately 12%."
    )

    document.add_paragraph(
        "Potential qualitative explanation: customer credit terms "
        "were extended from 30 days to 60 days for large "
        "multi-year contracts."
    )

    document.add_heading("Relevant Context Keywords", level=2)

    keywords = [
        "REL_01",
        "Accounts Receivable",
        "trade receivables",
        "Revenue",
        "credit terms",
        "customer payment terms",
        "30 days",
        "60 days",
        "extended payment terms",
        "multi-year contracts",
        "large customer contracts",
        "working capital",
        "customer collections",
        "outstanding invoices",
        "credit exposure",
    ]

    document.add_paragraph(", ".join(keywords))

    document.add_heading("Management Conclusion", level=2)

    document.add_paragraph(
        "Management believes that the increase in Accounts "
        "Receivable relative to Revenue is primarily explained "
        "by the temporary impact of extended customer payment "
        "terms on selected large multi-year contracts."
    )

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

    document.save(OUTPUT_FILE)

    print("SUCCESS")
    print(f"Created: {OUTPUT_FILE}")


if __name__ == "__main__":
    create_document()
